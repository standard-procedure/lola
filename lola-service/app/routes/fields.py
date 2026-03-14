"""GET /fields — Extract MERGEFIELD names from a .docx template using python-docx XML parsing."""

import os
import re

from fastapi.responses import JSONResponse

from app.config import config


def resolve_path(relative_path: str) -> str:
    """Resolve a path relative to the documents volume."""
    return os.path.join(config.DOCUMENTS_PATH, relative_path)


def extract_merge_fields(docx_path: str) -> list[str]:
    """
    Parse MERGEFIELD codes from docx XML using python-docx.

    Handles both simple fields (w:fldSimple) and complex fields
    (w:instrText between w:fldChar begin/end), including split instrText runs.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    fields: set[str] = set()

    # Helper to search paragraphs recursively (body + headers + footers)
    def scan_paragraphs(container):
        for para in container.paragraphs:
            scan_paragraph(para)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_paragraphs(cell)

    def scan_paragraph(para):
        # Collect all instrText segments between begin/end markers
        instr_parts = []
        in_field = 0  # nesting depth

        for child in para._p.iter():
            tag = child.tag

            if tag == qn("w:fldChar"):
                char_type = child.get(qn("w:fldCharType"), "")
                if char_type == "begin":
                    in_field += 1
                    instr_parts = [] if in_field == 1 else instr_parts
                elif char_type == "end":
                    if in_field == 1 and instr_parts:
                        _parse_instr("".join(instr_parts), fields)
                        instr_parts = []
                    in_field = max(in_field - 1, 0)

            elif tag == qn("w:instrText") and in_field > 0:
                instr_parts.append(child.text or "")

            elif tag == qn("w:fldSimple"):
                instr = child.get(qn("w:instr"), "")
                _parse_instr(instr, fields)

    def _parse_instr(instr: str, out: set):
        m = re.search(r"MERGEFIELD\s+([^\s\\]+)", instr, re.IGNORECASE)
        if m:
            out.add(m.group(1).strip('"'))

    # Scan document body
    scan_paragraphs(doc)

    # Scan headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                scan_paragraphs(hf)

    return sorted(fields)


async def handle_fields(template_path: str):
    from app.main import FieldsResponse

    abs_path = resolve_path(template_path)

    if not os.path.exists(abs_path):
        return JSONResponse(
            status_code=404,
            content={
                "error": f"File not found: {template_path}",
                "code": "TEMPLATE_NOT_FOUND",
            },
        )

    try:
        field_names = extract_merge_fields(abs_path)
    except Exception as e:
        return JSONResponse(
            status_code=422,
            content={
                "error": f"Cannot open template: {e}",
                "code": "TEMPLATE_ERROR",
            },
        )

    return FieldsResponse(
        template_path=template_path,
        fields=field_names,
        field_count=len(field_names),
    )
