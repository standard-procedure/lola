"""Tests for the POST /mail_merge endpoint."""

import os
import pytest
from unittest.mock import MagicMock, patch
from fastapi.testclient import TestClient
from app.main import app, get_uno_client


def make_client(merge_result=None, merge_side_effect=None):
    mock = MagicMock()
    mock.is_connected.return_value = True
    if merge_side_effect:
        mock.mail_merge.side_effect = merge_side_effect
    else:
        mock.mail_merge.return_value = merge_result or []
    app.dependency_overrides[get_uno_client] = lambda: mock
    return TestClient(app, raise_server_exceptions=False)


def _create_merge_template(path, field_names):
    """Create a .docx template with MERGEFIELD codes."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    para = doc.add_paragraph()
    for field_name in field_names:
        run_begin = para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_begin)

        run_instr = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = f" MERGEFIELD {field_name} \\* MERGEFORMAT "
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_instr._r.append(instr)

        run_sep = para.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)

        para.add_run(f"«{field_name}»")

        run_end = para.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)

    doc.save(str(path))


def test_mail_merge_returns_output_files(tmp_path):
    """POST /mail_merge executes merge and returns list of output file paths."""
    template = tmp_path / "letter.docx"
    _create_merge_template(template, ["CustomerName", "City"])
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    fake_output = output_dir / "0001.pdf"
    fake_output.write_bytes(b"pdf")

    client = make_client(merge_result=[str(fake_output)])

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "letter.docx",
            "data": [{"CustomerName": "Alice", "City": "London"}],
            "output_dir": "output",
            "output_format": "pdf",
        })

    assert response.status_code == 200
    data = response.json()
    assert data["record_count"] == 1
    assert len(data["output_files"]) == 1
    assert "duration_ms" in data
    assert data["warnings"] == []


def test_mail_merge_returns_400_for_empty_data(tmp_path):
    """POST /mail_merge returns 400 when data array is empty."""
    template = tmp_path / "letter.docx"
    _create_merge_template(template, ["Name"])
    client = make_client()

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "letter.docx",
            "data": [],
            "output_dir": "output",
        })

    assert response.status_code == 400
    data = response.json()
    assert data["code"] == "INVALID_REQUEST"


def test_mail_merge_returns_404_for_missing_template(tmp_path):
    """POST /mail_merge returns 404 when template does not exist."""
    client = make_client()

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "missing.docx",
            "data": [{"Name": "Alice"}],
            "output_dir": "output",
        })

    assert response.status_code == 404
    data = response.json()
    assert data["code"] == "TEMPLATE_NOT_FOUND"


def test_mail_merge_returns_400_for_too_many_records(tmp_path):
    """POST /mail_merge returns 400 when data exceeds 1000 records."""
    template = tmp_path / "letter.docx"
    _create_merge_template(template, ["Name"])
    client = make_client()
    data_records = [{"Name": f"Person {i}"} for i in range(1001)]

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "letter.docx",
            "data": data_records,
            "output_dir": "output",
        })

    assert response.status_code == 400
    data = response.json()
    assert data["code"] == "INVALID_REQUEST"


def test_mail_merge_returns_422_on_merge_failure(tmp_path):
    """POST /mail_merge returns 422 when LibreOffice merge fails."""
    template = tmp_path / "letter.docx"
    _create_merge_template(template, ["Name"])
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    from app.exceptions import MergeError
    client = make_client(merge_side_effect=MergeError("LibreOffice could not process the template"))

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "letter.docx",
            "data": [{"Name": "Alice"}],
            "output_dir": "output",
        })

    assert response.status_code == 422
    data = response.json()
    assert data["code"] == "MERGE_ERROR"


def test_mail_merge_returns_503_when_lo_unavailable(tmp_path):
    """POST /mail_merge returns 503 when LibreOffice is not connected."""
    template = tmp_path / "letter.docx"
    _create_merge_template(template, ["Name"])
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    from app.exceptions import LibreOfficeError
    client = make_client(merge_side_effect=LibreOfficeError("LibreOffice is not available"))

    with patch("app.routes.mail_merge.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.post("/mail_merge", json={
            "template_path": "letter.docx",
            "data": [{"Name": "Alice"}],
            "output_dir": "output",
        })

    assert response.status_code == 503
    data = response.json()
    assert data["code"] == "LIBREOFFICE_ERROR"
