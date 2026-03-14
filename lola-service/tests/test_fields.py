"""Tests for the GET /fields endpoint."""

import os
import pytest
from unittest.mock import patch
from fastapi.testclient import TestClient
from app.main import app, get_uno_client
from unittest.mock import MagicMock


def make_client():
    mock = MagicMock()
    mock.is_connected.return_value = True
    app.dependency_overrides[get_uno_client] = lambda: mock
    return TestClient(app, raise_server_exceptions=False)


def _create_docx_with_fields(path, field_names):
    """Create a minimal .docx file with MERGEFIELD codes in the XML."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    doc = Document()
    para = doc.add_paragraph()

    for field_name in field_names:
        # Add fldChar begin
        run_begin = para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_begin)

        # Add instrText
        run_instr = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = f" MERGEFIELD {field_name} \\* MERGEFORMAT "
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_instr._r.append(instr)

        # Add fldChar separate
        run_sep = para.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)

        # Add display text
        para.add_run(f"«{field_name}»")

        # Add fldChar end
        run_end = para.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)

    doc.save(str(path))


def test_fields_extracts_merge_fields(tmp_path):
    """GET /fields returns sorted, deduplicated merge field names."""
    template = tmp_path / "letter.docx"
    _create_docx_with_fields(template, ["City", "CustomerName", "Address", "City"])  # City duplicated

    client = make_client()

    with patch("app.routes.fields.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.get("/fields", params={"template_path": "letter.docx"})

    assert response.status_code == 200
    data = response.json()
    assert data["fields"] == ["Address", "City", "CustomerName"]
    assert data["field_count"] == 3
    assert data["template_path"] == "letter.docx"


def test_fields_returns_empty_for_document_with_no_fields(tmp_path):
    """GET /fields returns empty list for a document with no merge fields."""
    from docx import Document
    template = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.save(str(template))

    client = make_client()

    with patch("app.routes.fields.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.get("/fields", params={"template_path": "plain.docx"})

    assert response.status_code == 200
    data = response.json()
    assert data["fields"] == []
    assert data["field_count"] == 0


def test_fields_returns_404_for_missing_template(tmp_path):
    """GET /fields returns 404 when template file does not exist."""
    client = make_client()

    with patch("app.routes.fields.resolve_path", side_effect=lambda p: str(tmp_path / p)):
        response = client.get("/fields", params={"template_path": "missing.docx"})

    assert response.status_code == 404
    data = response.json()
    assert data["code"] == "TEMPLATE_NOT_FOUND"
